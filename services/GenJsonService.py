import os
import re
import json
import docx
from difflib import SequenceMatcher
from dotenv import load_dotenv
from typing import List
from google import genai
from services.types.Question import Question
import openai 

class GenJsonService:
    def __init__(self):
        load_dotenv()
        self.prompt1 = """You are given two texts extracted from a DOCX:
    1) Full text with questions and choices.
    2) Colored text containing correct answers.
    Generate a JSON array of objects with fields [question, choices, answer].
    Ensure each answer is in its choices. If the choices/answer don't start with ก. ข. ค. ง., add these as prefixes.
    Output a concise textual representation, then produce valid JSON data."""
        self.geminiLlm = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))

    def readDocText(self, path):
        doc = docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    def readDocColoredText(self, path):
        doc = docx.Document(path)
        colored = ""
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.color.rgb or run.font.highlight_color:
                    colored += run.text
        return colored.strip()

    def listFiles(self, directory):
        return os.listdir(directory)

    def reformat(self, text: str):
        text = text.strip()
        text = text[3:] if text.startswith("```") else text
        text = text[:-3] if text.endswith("```") else text
        text = text.replace('\n', ' ').strip()
        return re.sub(r'\s+', ' ', text).strip()

    def removeNewlines(self, data):
        if isinstance(data, str):
            return data.replace("\n", "")
        if isinstance(data, list):
            return [self.removeNewlines(i) for i in data]
        if isinstance(data, dict):
            return {k: self.removeNewlines(v) for k, v in data.items()}
        return data

    def getMostSimilarChoice(self, choices: List[str], answer: str):
        return max(choices, key=lambda c: SequenceMatcher(None, c, answer).ratio())

    def checker(self, questions: List[Question]):
        qNoFour, ansNotIn = [], []
        for i, q in enumerate(questions, start=1):
            if q['answer'] not in q['choices']:
                ansNotIn.append(i)
            if len(q['choices']) != 4:
                qNoFour.append(i)
        return len(questions), qNoFour, ansNotIn

    def processFile(self, filePath):
        allText = self.readDocText(filePath)
        colorText = self.readDocColoredText(filePath)
        messages = [
            ("system", self.prompt1),
            ("human", f"All Text: {allText}\n\nColor Text: {colorText}")
        ]
        # Save the request to llm_req.txt file
        with open('tmp/llm_req.txt', 'w', encoding='utf-8') as f:
            f.write(f"System: {self.prompt1}\n\n")
            f.write(f"Human: All Text: {allText}\n\nColor Text: {colorText}")
        
        response = self.geminiLlm.models.generate_content(
            model='gemini-2.5-pro',
            contents=messages,
            config={
                'response_mime_type': 'application/json',
                'response_schema': list[Question]
            }
        )
        
        # Save the response to tmp/llm_result.json
        with open('tmp/llm_result.json', 'w', encoding='utf-8') as f:
            response_data = json.loads(response.text)
            json.dump(response_data, f, ensure_ascii=False, indent=2)

        input("Press Enter after saving the response to tmp/llm_result.json file...")
        
        with open('tmp/llm_result.json', 'r', encoding='utf-8') as f:
            response = f.read()
        finalOutput = self.reformat(response)
        finalOutput = json.loads(finalOutput)
        for q in finalOutput:
            q['answer'] = self.getMostSimilarChoice(q['choices'], q['answer'])
        questionCount, noFour, ansNotIn = self.checker(finalOutput)
        print(f"Number of questions: {questionCount}")
        print(f"Questions without four choices: {noFour}")
        print(f"Answers not in choices: {ansNotIn}")
        return finalOutput

    def run(self, inputDir='./input', outputDir='./tmp', outputFile='output.json'):
        print("Step 1: Listing files in directory:", inputDir)
        files = self.listFiles(inputDir)
        if not files:
            print("No files found.")
            return
        firstFile = os.path.join(inputDir, files[0])
        print("Step 2: Processing file:", firstFile)
        data = self.processFile(firstFile)
        print("Step 3: Removing newlines from data")
        data = self.removeNewlines(data)
        if not os.path.exists(outputDir):
            print("Step Err: Creating output directory:", outputDir)
            os.makedirs(outputDir)
        outputFile = os.path.join(outputDir, outputFile)
        print("Step 4: Writing output to file:", outputFile)
        with open(outputFile, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
